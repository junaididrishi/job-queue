"""
Generate Job_Queue_Complete_Guide.docx — Python/FastAPI edition (detailed).
Run: python3 scripts/make_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Colour palette ───────────────────────────────────────────────────────────
INK      = RGBColor(0x1E, 0x29, 0x3B)
BLUE     = RGBColor(0x01, 0x5F, 0xA3)
DARKBLUE = RGBColor(0x1E, 0x40, 0xAF)
GREEN    = RGBColor(0x0F, 0x5A, 0x2E)
GRAY     = RGBColor(0x47, 0x55, 0x69)
BROWN    = RGBColor(0x78, 0x35, 0x00)
PURPLE   = RGBColor(0x5B, 0x21, 0xB6)
YELLOW_BG = "FFFBEB"
CODE_BG   = "F1F5F9"
QA_BG     = "EFF6FF"
WARN_BG   = "FEF2F2"
TIP_BG    = "F0FDF4"


def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22 if level == 1 else 14)
    p.paragraph_format.space_after  = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = BLUE
    run.font.size = Pt(18 if level == 1 else 14)
    return p


def subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = PURPLE
    run.font.size = Pt(12)
    return p


def body(doc, text, bold=False, italic=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent * 0.8)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.color.rgb = INK
    run.font.size = Pt(11)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.6)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.color.rgb = INK
    run.font.size = Pt(10.5)
    return p


def numbered(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.6)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.color.rgb = INK
    run.font.size = Pt(10.5)
    return p


def code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_bg(cell, CODE_BG)
    cell.paragraphs[0].clear()
    for line in code_text.strip().split("\n"):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = GREEN
    doc.add_paragraph()


def note_box(doc, label, text, bg=YELLOW_BG, text_color=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_bg(cell, bg)
    cell.paragraphs[0].clear()
    p = cell.add_paragraph()
    run_label = p.add_run(label + "  ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = BROWN if bg == YELLOW_BG else RGBColor(0x06, 0x60, 0x27)
    run_body = p.add_run(text)
    run_body.font.size = Pt(10)
    run_body.font.color.rgb = text_color or INK
    doc.add_paragraph()


def qa_block(doc, q, a):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_bg(cell, QA_BG)
    cell.paragraphs[0].clear()

    pq = cell.add_paragraph()
    rq = pq.add_run("Q:  " + q)
    rq.bold = True
    rq.font.size = Pt(11)
    rq.font.color.rgb = DARKBLUE

    pa = cell.add_paragraph()
    ra = pa.add_run("A:  " + a)
    ra.font.size = Pt(10.5)
    ra.font.color.rgb = INK
    doc.add_paragraph()


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 90)
    run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    run.font.size = Pt(8)


# ─────────────────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(36)
title_run = title_p.add_run("Distributed Job Queue System")
title_run.bold = True
title_run.font.size = Pt(28)
title_run.font.color.rgb = BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("Complete Beginner Guide  ·  Tech Stack Deep Dive  ·  Interview Q&A")
sub_run.font.size = Pt(13)
sub_run.font.color.rgb = GRAY

stack_p = doc.add_paragraph()
stack_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
stack_run = stack_p.add_run(
    "Python  ·  FastAPI  ·  Redis  ·  PostgreSQL  ·  Docker  ·  AWS EC2  ·  GitHub Actions  ·  Locust  ·  Sentry"
)
stack_run.font.size = Pt(10.5)
stack_run.font.color.rgb = GRAY
doc.add_paragraph()
divider(doc)
doc.add_paragraph()

toc_intro = doc.add_paragraph()
toc_run = toc_intro.add_run("Table of Contents")
toc_run.bold = True
toc_run.font.size = Pt(13)
toc_run.font.color.rgb = BLUE

toc_items = [
    "1.  The Problem — Why Job Queues Exist",
    "2.  Core Concepts — What Every Beginner Must Understand",
    "3.  End-to-End Flow — How This System Works",
    "4.  Tech Stack — Every Tool Explained in Depth",
    "5.  Database Schema — Tables and Their Purpose",
    "6.  Project Structure — Every File and Why It Exists",
    "7.  Environment Variables — Configuration Guide",
    "8.  Running Locally — Step-by-Step Setup",
    "9.  Running Tests — Unit, Integration, and Load Tests",
    "10. CI/CD Pipeline — Automated Build and Deploy",
    "11. Deploying to AWS EC2 — Production Setup",
    "12. Security — Auth, Webhooks, Secrets",
    "13. Observability — Logging, Sentry, Dashboard",
    "14. Load Testing — Results and What We Fixed",
    "15. Scaling the System — What Changes at 10x Load",
    "16. Common Mistakes and How to Avoid Them",
    "17. Interview Questions & Detailed Answers (20 Q&As)",
    "18. Glossary",
]
for item in toc_items:
    bullet(doc, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — THE PROBLEM
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "1. The Problem — Why Job Queues Exist")

body(doc, (
    "Imagine you are building a web app. A user clicks 'Sign Up'. Your API needs to: "
    "save their account to the database, send them a welcome email, resize their profile picture, "
    "and notify your analytics service. If you do all of this inside the HTTP request handler, "
    "the user waits 3-5 seconds before they get a response. That is a terrible user experience."
))

body(doc, (
    "The real problem is that not all work needs to happen before you respond. "
    "Saving the account is required immediately — the user needs their ID. "
    "But sending the email, resizing the photo, and notifying analytics can happen "
    "one second later without the user noticing. A job queue lets you separate "
    "'work that must happen now' from 'work that can happen soon'."
))

subheading(doc, "The Restaurant Analogy")
body(doc, (
    "Think of your API like a waiter and your worker like the kitchen. "
    "When you order food, the waiter does not cook it at your table — "
    "they write your order on a ticket, give it to the kitchen, and immediately go serve other customers. "
    "The kitchen processes tickets in order. "
    "You (the customer) get acknowledgement instantly and your food arrives later. "
    "The job queue is the ticket system between the waiter and the kitchen."
))

subheading(doc, "Without a Job Queue vs. With a Job Queue")
body(doc, "WITHOUT a job queue:", bold=True)
for item in [
    "POST /register takes 3-5 seconds (blocks while sending email)",
    "If the email server is slow or down, the request fails entirely",
    "You cannot retry failed operations automatically",
    "Under load, slow jobs consume all server threads — other requests queue up",
    "No visibility into what is being processed or why something failed",
]:
    bullet(doc, item)

body(doc, "WITH a job queue:", bold=True)
for item in [
    "POST /register returns in ~20ms (just saves to DB, pushes to queue)",
    "Email failure does not fail the registration — it retries automatically",
    "Exponential backoff prevents hammering a struggling mail server",
    "Workers scale independently from the API — add more workers for more throughput",
    "Full audit trail: every task has a status, logs, and retry count in the database",
]:
    bullet(doc, item)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — CORE CONCEPTS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "2. Core Concepts — What Every Beginner Must Understand")

subheading(doc, "2.1  Producer / Consumer Pattern")
body(doc, (
    "This is the foundational pattern of all job queues. A Producer creates work "
    "and puts it on a queue. A Consumer takes work off the queue and executes it. "
    "In our system, the FastAPI app is the Producer (it calls LPUSH to push task IDs onto Redis lists). "
    "The worker process is the Consumer (it calls BRPOP to pull task IDs off Redis lists). "
    "The two sides are decoupled — the producer does not care when or how the consumer processes the work."
))

subheading(doc, "2.2  LPUSH and BRPOP — Redis Queue Mechanics")
body(doc, (
    "Redis lists are the backbone of this queue. LPUSH adds an item to the LEFT (front) of a list. "
    "BRPOP removes and returns an item from the RIGHT (back) of a list — it is a Blocking Right POP. "
    "This gives us FIFO order: first item pushed is the first item consumed, "
    "because push goes to the left end and pop takes from the right end."
))
body(doc, (
    "The 'B' in BRPOP stands for Blocking. When the list is empty, BRPOP does not return immediately "
    "— it waits (blocks) until an item appears or a timeout is reached. "
    "This means worker threads sleep efficiently inside BRPOP rather than spinning in a loop checking "
    "every millisecond. The OS wakes them up the moment Redis has new work. No CPU waste."
))
note_box(doc, "KEY POINT:", (
    "BRPOP [tasks:high, tasks:normal, tasks:low] checks the lists IN ORDER. "
    "If tasks:high has items, they are always returned before tasks:normal items. "
    "Only when tasks:high is empty does Redis look at tasks:normal. "
    "This is how priority queuing is implemented — no custom code needed."
))

subheading(doc, "2.3  Exponential Backoff")
body(doc, (
    "When a job fails, you should not retry it immediately. Imagine your email server is overloaded — "
    "retrying immediately makes it worse. Exponential backoff means each retry waits longer than the last. "
    "In this system: delay = min(5000ms * 2^attempt, 3,600,000ms). "
    "Attempt 1 waits 10 seconds. Attempt 2 waits 20 seconds. Attempt 3 waits 40 seconds. "
    "The cap at 3,600,000ms (1 hour) prevents retries from waiting so long they become useless."
))

subheading(doc, "2.4  Dead-Letter Queue (DLQ)")
body(doc, (
    "A dead-letter queue is where tasks go when they have permanently failed — after all retries are exhausted. "
    "In our system, max_retries = 3. After the 3rd failure, instead of deleting the task, "
    "we push it to the tasks:dead Redis list and set its status to 'dead' in PostgreSQL. "
    "This is critical for observability: without a DLQ, failed tasks just disappear. "
    "With a DLQ, an operator can log in, inspect what failed and why, fix the underlying problem, "
    "and re-enqueue the task."
))

subheading(doc, "2.5  Idempotency")
body(doc, (
    "A job is idempotent if running it twice produces the same result as running it once. "
    "Job queues should aim for idempotent tasks because at-least-once delivery guarantees "
    "that a task might be processed more than once (e.g., worker crashes after completing the job "
    "but before acknowledging it). In this project, each task has a unique UUID. "
    "Handlers could be made idempotent by checking if the result already exists before doing work."
))

subheading(doc, "2.6  At-Least-Once vs. Exactly-Once Delivery")
body(doc, (
    "Redis BRPOP gives at-least-once delivery: a message is removed from the queue when BRPOP returns it. "
    "If the worker crashes after popping but before completing the job, the task is lost from the queue "
    "(though it remains in PostgreSQL as 'processing'). "
    "Exactly-once delivery is much harder and requires distributed transactions or acknowledgement patterns. "
    "For most use cases (emails, resizes, reports), at-least-once with idempotent handlers is sufficient."
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — END-TO-END FLOW
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "3. End-to-End Flow — How This System Works")

body(doc, (
    "Let us trace a single task — sending a welcome email — from the moment the client makes "
    "the API request to the moment the email is 'sent' and the webhook is fired."
))

subheading(doc, "Step 1: Client sends the request")
body(doc, "The client sends:")
code_block(doc, """\
POST /api/tasks
Authorization: Bearer eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...
Content-Type: application/json

{
  "type": "email",
  "priority": "high",
  "payload": {
    "to": "newuser@example.com",
    "subject": "Welcome to our platform!"
  },
  "webhook_url": "https://myapp.com/webhooks/job-complete"
}""")
body(doc, (
    "The Authorization header carries a JWT (JSON Web Token). "
    "FastAPI's dependency injection system extracts the token and verifies it before "
    "the route handler runs. If the token is missing or invalid, the request is rejected with 403."
))

subheading(doc, "Step 2: API validates and saves to PostgreSQL")
body(doc, (
    "Pydantic validates the request body. It checks that 'type' is one of the allowed enum values "
    "(email, image_resize, report_generation, data_export, notification), "
    "that 'priority' is high/normal/low, and that 'payload' is a valid JSON object. "
    "If validation fails, FastAPI automatically returns a 422 Unprocessable Entity with details about "
    "which field failed and why — no manual error handling needed."
))
body(doc, "A new row is inserted into the tasks table:")
code_block(doc, """\
INSERT INTO tasks (id, type, status, priority, payload, webhook_url, created_by, created_at)
VALUES (
  'a3f2b1c4-...',       -- UUID generated by the API
  'email',
  'pending',            -- initial status
  'high',
  '{"to": "newuser@example.com", "subject": "..."}',
  'https://myapp.com/webhooks/job-complete',
  42,                   -- user ID from the JWT
  NOW()
);""")
body(doc, (
    "The task is saved to PostgreSQL first, before touching Redis. "
    "This ensures the task is never lost — even if Redis crashes, the task record exists in the database. "
    "A background monitoring job could scan for 'pending' tasks with no Redis entry and re-enqueue them."
))

subheading(doc, "Step 3: API pushes to Redis queue")
body(doc, (
    "After saving to PostgreSQL, the API calls LPUSH on the appropriate Redis list based on priority:"
))
code_block(doc, """\
# Python code in app/queue/queue.py
await redis.lpush("tasks:high", json.dumps({
    "task_id": "a3f2b1c4-...",
    "task_type": "email",
    "priority": "high"
}))""")
body(doc, (
    "The message contains the task ID and type — not the full payload. "
    "The worker will load the full task (including payload and webhook URL) from PostgreSQL. "
    "This avoids putting large payloads in Redis, keeps Redis memory usage low, "
    "and ensures the worker always has the latest task state."
))
body(doc, "The API immediately returns 202 Accepted:")
code_block(doc, """\
HTTP/1.1 202 Accepted
{
  "id": "a3f2b1c4-...",
  "type": "email",
  "status": "pending",
  "priority": "high",
  "created_at": "2026-07-06T10:23:45Z"
}""")
body(doc, (
    "202 (not 200) signals to the client: 'We accepted your request, but it has not been processed yet.' "
    "This is the correct HTTP status for async operations."
))

subheading(doc, "Step 4: Worker dequeues the task")
body(doc, (
    "The worker process runs independently, in a separate process from the API. "
    "It has a ThreadPoolExecutor with N threads (default: 4). Each thread runs a loop:"
))
code_block(doc, """\
# Simplified from app/worker/worker.py
def worker_thread(thread_id):
    while not _shutdown.is_set():
        # BRPOP blocks here until a task arrives (or timeout)
        result = redis.brpop(["tasks:high", "tasks:normal", "tasks:low"], timeout=5)
        if result is None:
            continue  # timeout — check shutdown flag, then block again

        queue_name, raw_message = result
        message = json.loads(raw_message)
        process_one(message)""")
body(doc, (
    "BRPOP checks tasks:high first. If it has items, one is returned immediately. "
    "If tasks:high is empty, it checks tasks:normal. If that is also empty, it checks tasks:low. "
    "If all three are empty, the call blocks for up to 5 seconds (the timeout), then returns None. "
    "The thread checks the shutdown flag and blocks again. "
    "This means worker threads use zero CPU when there is no work to do."
))

subheading(doc, "Step 5: Worker loads full task and executes the job")
body(doc, "The worker fetches the full task from PostgreSQL, updates status to 'processing', and calls the handler:")
code_block(doc, """\
# In process_one():
task = get_task_sync(session, task_id)       # load from PostgreSQL
update_task_sync(session, task_id, status="processing", started_at=now)

try:
    result = execute_job(task.type, task.payload)  # runs the handler
    update_task_sync(session, task_id, status="completed", result=result)
    if task.webhook_url:
        deliver_webhook(task)                  # fire callback
except Exception as e:
    if task.retry_count < MAX_RETRIES:
        delay = retry_delay_ms(task.retry_count + 1)
        update_task_sync(session, task_id, status="retrying", retry_count=task.retry_count+1)
        threading.Timer(delay / 1000, re_enqueue, args=[task]).start()
    else:
        update_task_sync(session, task_id, status="dead")
        enqueue_dead_letter_sync(task_id, str(e))""")

subheading(doc, "Step 6: Webhook callback")
body(doc, (
    "If the task had a webhook_url, the worker sends a POST request to it with the result. "
    "The request is HMAC-SHA256 signed to prove it came from this system (explained in Section 12). "
    "The receiver can verify the signature and trust the payload."
))
code_block(doc, """\
POST https://myapp.com/webhooks/job-complete
X-JobQueue-Signature: sha256=a4f3c2b1e9...
Content-Type: application/json

{
  "task_id": "a3f2b1c4-...",
  "type": "email",
  "status": "completed",
  "result": {"message_id": "msg_789", "provider": "simulated-smtp"},
  "completed_at": "2026-07-06T10:23:46Z"
}""")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — TECH STACK
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "4. Tech Stack — Every Tool Explained in Depth")

subheading(doc, "4.1  Python 3.12")
body(doc, (
    "Python is the language. Version 3.12 brought significant performance improvements "
    "(up to 60% faster than 3.10 on benchmarks) and better error messages. "
    "We use Python's asyncio for the API (concurrent I/O without threads) "
    "and Python's threading module for the worker (blocking I/O that cannot use asyncio)."
))

subheading(doc, "4.2  FastAPI")
body(doc, (
    "FastAPI is a modern Python web framework built on top of Starlette (the async HTTP toolkit) "
    "and Pydantic (data validation). Key features used in this project:"
))
for feat in [
    "async route handlers — handles many concurrent requests without blocking",
    "Pydantic models as function parameters — automatic request validation with clear error messages",
    "Dependency Injection (Depends) — get_db() and get_current_user() injected cleanly",
    "OpenAPI auto-docs — visit /docs for a live Swagger UI with all endpoints, no manual work",
    "HTTPException — raises proper HTTP error responses with a single line",
    "Lifespan context manager — runs startup/shutdown code (connect to Redis, close connections)",
]:
    bullet(doc, feat)
body(doc, (
    "Why not Flask? Flask is synchronous by default. Adding async to Flask requires extra libraries "
    "and is not first-class. Why not Django? Django has a full ORM, templating, and admin panel — "
    "90% of which we do not need. FastAPI gives us exactly what a JSON API needs and nothing more."
))

subheading(doc, "4.3  PostgreSQL")
body(doc, (
    "PostgreSQL is the relational database that stores all persistent data: "
    "tasks, users, workers, and task logs. "
    "We chose PostgreSQL over other databases for these reasons:"
))
for reason in [
    "ACID transactions: task creation and status updates are atomic — no partial writes",
    "JSONB column type: task payloads and results are stored as JSON, queryable with indexes",
    "Rich query language: the stats endpoint uses COUNT(*) FILTER (WHERE status='pending') for efficient aggregation",
    "pg_isready health check: Docker Compose waits for PostgreSQL to be truly ready, not just running",
]:
    bullet(doc, reason)
body(doc, (
    "We use two different SQLAlchemy drivers: asyncpg for the API (async, high performance) "
    "and psycopg2 for the worker (sync, required for threading). "
    "This is why the .env has two separate DATABASE_URL variables."
))

subheading(doc, "4.4  Redis")
body(doc, (
    "Redis is an in-memory data structure server. We use it as the message broker "
    "— the channel between the API (producer) and the worker (consumer). "
    "Redis is ideal for this role because:"
))
for reason in [
    "Extremely fast: sub-millisecond LPUSH/BRPOP operations",
    "BRPOP blocks efficiently: workers sleep in the kernel, not in a polling loop",
    "Ordered lists: FIFO queue built-in, no custom queue implementation needed",
    "Multiple named lists: tasks:high, tasks:normal, tasks:low — priority built into the data structure",
    "Pub/Sub and Streams available if we want to upgrade later",
]:
    bullet(doc, reason)
note_box(doc, "IMPORTANT:", (
    "Redis is not a database. If Redis crashes, any messages in the queue that have not been "
    "popped yet are lost (unless you enable Redis persistence with AOF). "
    "This is acceptable because: (a) tasks are saved to PostgreSQL first, and "
    "(b) a recovery job could scan for 'pending' tasks older than N minutes and re-enqueue them."
))

subheading(doc, "4.5  SQLAlchemy 2.0 and Alembic")
body(doc, (
    "SQLAlchemy is the ORM (Object-Relational Mapper) — it lets us define database tables as Python classes "
    "and write queries in Python instead of raw SQL. "
    "Version 2.0 introduced a cleaner async API using async with session.begin()."
))
body(doc, (
    "Alembic is the database migration tool. Instead of running ALTER TABLE commands manually, "
    "you write migration scripts in Python. Each migration has an upgrade() function "
    "(apply the change) and a downgrade() function (reverse it). "
    "When the app starts in production, alembic upgrade head runs all pending migrations "
    "to bring the database schema up to date."
))
code_block(doc, """\
# How a migration looks (alembic/versions/001_initial_schema.py)
def upgrade():
    op.create_table(
        'tasks',
        sa.Column('id', sa.String(36), primary_key=True),
        sa.Column('type', sa.String(50), nullable=False),
        sa.Column('status', sa.String(20), nullable=False, server_default='pending'),
        # ... more columns
    )
    op.create_index('ix_tasks_status', 'tasks', ['status'])

def downgrade():
    op.drop_table('tasks')""")

subheading(doc, "4.6  Docker and Docker Compose")
body(doc, (
    "Docker packages the application and all its dependencies into an image. "
    "When you run that image anywhere (your laptop, a CI server, AWS EC2), "
    "it behaves identically. No more 'works on my machine' problems."
))
body(doc, "Our Dockerfile:")
code_block(doc, """\
FROM python:3.12-slim

# Install system libs required by psycopg2 (C extension for PostgreSQL)
RUN apt-get update && apt-get install -y gcc libpq-dev && rm -rf /var/lib/apt/lists/*

WORKDIR /app
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt
COPY . .

HEALTHCHECK --interval=30s --timeout=10s --start-period=20s --retries=3 \\
  CMD curl -f http://localhost:8000/health || exit 1

CMD ["uvicorn", "app.main:app", "--host", "0.0.0.0", "--port", "8000"]""")
body(doc, (
    "The same image is used for both the API and the Worker. "
    "The only difference is the CMD. In docker-compose.yml, the worker service overrides it:"
))
code_block(doc, """\
# docker-compose.yml (simplified)
services:
  postgres:
    image: postgres:16-alpine
    environment: {POSTGRES_DB: job_queue, POSTGRES_USER: postgres, POSTGRES_PASSWORD: password}
    healthcheck:
      test: [CMD, pg_isready, -U, postgres]

  redis:
    image: redis:7-alpine
    healthcheck:
      test: [CMD, redis-cli, ping]

  api:
    build: .
    ports: [8000:8000]
    depends_on:
      postgres: {condition: service_healthy}
      redis:    {condition: service_healthy}
    env_file: .env

  worker:
    build: .
    command: python -m app.worker.worker   # override the default CMD
    depends_on:
      api: {condition: service_started}
    deploy:
      replicas: 1    # increase to 3 to scale out""")

subheading(doc, "4.7  pytest — Testing Framework")
body(doc, (
    "pytest is the standard Python testing framework. "
    "We have two test layers in this project:"
))
body(doc, "Unit tests (tests/unit/):", bold=True)
for item in [
    "Test individual functions in isolation",
    "No database or Redis connection needed",
    "Run in milliseconds — fast feedback during development",
    "Examples: retry_delay_ms() math, password hashing, JWT encode/decode, job handler logic",
]:
    bullet(doc, item)
body(doc, "Integration tests (tests/integration/):", bold=True)
for item in [
    "Test the full API using httpx AsyncClient with ASGITransport",
    "ASGITransport sends requests directly to the FastAPI app, no network socket needed",
    "Require a real PostgreSQL and Redis (provided by GitHub Actions service containers in CI)",
    "Examples: register user, login, create task, check status, 404 on missing task",
]:
    bullet(doc, item)

subheading(doc, "4.8  GitHub Actions (CI/CD)")
body(doc, (
    "GitHub Actions is GitHub's built-in automation system. "
    "You define workflows in YAML files. Workflows run on GitHub's servers (Ubuntu VMs) "
    "when events happen (push, pull request, schedule). "
    "We use it for a 4-stage pipeline: test → test with real services → build Docker images → deploy."
))

subheading(doc, "4.9  AWS EC2")
body(doc, (
    "EC2 (Elastic Compute Cloud) is a virtual machine on AWS. "
    "We run a single Ubuntu 22.04 instance (t3.small: 2 vCPUs, 2 GB RAM) "
    "and run all 4 Docker Compose services on it. "
    "For production scale, you would use: "
    "AWS RDS for PostgreSQL, AWS ElastiCache for Redis, "
    "and multiple EC2 instances behind an Application Load Balancer. "
    "But for a portfolio project, one EC2 instance is sufficient and easy to explain."
))

subheading(doc, "4.10  Locust — Load Testing")
body(doc, (
    "Locust is a Python load testing tool. You write load scenarios in Python: "
    "define user behavior (what endpoints to hit, how often, with what data), "
    "then Locust spawns thousands of simulated concurrent users. "
    "It reports request rates, latencies (p50, p95, p99), and error rates in real time."
))

subheading(doc, "4.11  structlog — Structured Logging")
body(doc, (
    "Regular logging produces text like: '2026-07-06 10:23:45 INFO Task completed'. "
    "Structured logging produces JSON: {'timestamp': '...', 'level': 'info', 'event': 'task_completed', "
    "'task_id': 'a3f2b1c4', 'duration_ms': 234}. "
    "JSON logs can be ingested by logging aggregators (CloudWatch, Datadog, Logtail) and "
    "queried/filtered programmatically. In development, structlog renders colored, readable console output."
))

subheading(doc, "4.12  Sentry — Error Monitoring")
body(doc, (
    "Sentry captures unhandled exceptions in real time and sends you an alert. "
    "For each error it records: the full stack trace, the request that caused it, "
    "the user (if logged in), environment variables (non-secret), and a count of how many times it occurred. "
    "It also does performance tracing: tracks how long each API endpoint takes end-to-end "
    "and identifies which DB query or Redis call is the bottleneck. "
    "We set traces_sample_rate=0.2 meaning 20% of requests are traced (100% would be too expensive)."
))

subheading(doc, "4.13  JWT — JSON Web Tokens")
body(doc, (
    "JWT is a standard for stateless authentication. "
    "When a user logs in, the API creates a signed token containing their user ID and email. "
    "The token is base64-encoded JSON: header.payload.signature. "
    "The signature is computed with the SECRET_KEY — only the API can create valid tokens. "
    "The client stores the token and sends it with every request in the Authorization header. "
    "The API decodes and verifies the token on each request — no database lookup needed. "
    "Tokens expire (configurable) to limit the window for stolen tokens."
))
code_block(doc, """\
# Decoded JWT payload looks like:
{
  "sub": "42",              # user ID
  "email": "user@example.com",
  "exp": 1782345678,        # expiry timestamp
  "iat": 1782259278         # issued-at timestamp
}""")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — DATABASE SCHEMA
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "5. Database Schema — Tables and Their Purpose")

subheading(doc, "5.1  users table")
code_block(doc, """\
CREATE TABLE users (
  id            SERIAL PRIMARY KEY,
  email         VARCHAR(255) UNIQUE NOT NULL,
  password_hash VARCHAR(255) NOT NULL,   -- bcrypt hash, never plain text
  name          VARCHAR(255),
  created_at    TIMESTAMPTZ DEFAULT NOW()
);""")
body(doc, (
    "Stores registered users. The password_hash column stores a bcrypt hash — "
    "bcrypt is a slow hashing algorithm specifically designed for passwords. "
    "It includes a random salt (prevents rainbow table attacks) and a cost factor "
    "(makes brute-force expensive). Even if the database leaks, passwords are safe."
))

subheading(doc, "5.2  tasks table")
code_block(doc, """\
CREATE TABLE tasks (
  id            VARCHAR(36)  PRIMARY KEY,  -- UUID
  type          VARCHAR(50)  NOT NULL,      -- email, image_resize, etc.
  status        VARCHAR(20)  NOT NULL DEFAULT 'pending',
                                            -- pending, processing, completed, retrying, dead
  priority      VARCHAR(10)  NOT NULL DEFAULT 'normal',  -- high, normal, low
  payload       JSONB        NOT NULL,      -- job-specific input data
  result        JSONB,                      -- job output (set on completion)
  error_message TEXT,                       -- last error (set on failure)
  retry_count   INTEGER      NOT NULL DEFAULT 0,
  max_retries   INTEGER      NOT NULL DEFAULT 3,
  webhook_url   TEXT,                       -- optional callback URL
  webhook_delivered BOOLEAN  DEFAULT FALSE,
  created_by    INTEGER REFERENCES users(id),
  created_at    TIMESTAMPTZ  DEFAULT NOW(),
  started_at    TIMESTAMPTZ,
  completed_at  TIMESTAMPTZ
);

-- Indexes for common query patterns:
CREATE INDEX ix_tasks_status          ON tasks(status);
CREATE INDEX ix_tasks_priority_status ON tasks(priority, status);
CREATE INDEX ix_tasks_created_at      ON tasks(created_at DESC);""")
body(doc, (
    "The three indexes are chosen based on the query patterns: "
    "GET /api/tasks filters by status, the stats endpoint groups by status, "
    "and list endpoints order by created_at descending (newest first)."
))

subheading(doc, "5.3  task_logs table")
code_block(doc, """\
CREATE TABLE task_logs (
  id         SERIAL PRIMARY KEY,
  task_id    VARCHAR(36) REFERENCES tasks(id) ON DELETE CASCADE,
  level      VARCHAR(10),    -- info, warning, error
  message    TEXT,
  metadata_  JSONB,           -- extra context (retry count, error type, etc.)
  created_at TIMESTAMPTZ DEFAULT NOW()
);""")
body(doc, (
    "Every state transition writes a log entry. This gives a full audit trail: "
    "when did processing start, how many retries, what was the error each time. "
    "GET /api/tasks/:id includes these logs in the response."
))

subheading(doc, "5.4  workers table")
code_block(doc, """\
CREATE TABLE workers (
  id               VARCHAR(36) PRIMARY KEY,  -- UUID assigned at startup
  hostname         VARCHAR(255),
  pid              INTEGER,
  status           VARCHAR(20),  -- active, idle, offline
  current_task_id  VARCHAR(36),
  tasks_processed  INTEGER DEFAULT 0,
  tasks_failed     INTEGER DEFAULT 0,
  last_heartbeat   TIMESTAMPTZ DEFAULT NOW()
);""")
body(doc, (
    "When the worker process starts, it registers itself in this table. "
    "Every 30 seconds it updates last_heartbeat. The dashboard shows all registered workers "
    "and their stats. If a worker's last_heartbeat is older than 60 seconds, it is considered offline."
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "6. Project Structure — Every File and Why It Exists")

code_block(doc, """\
job-queue/
├── app/
│   ├── __init__.py
│   ├── main.py                  FastAPI app creation, middleware, route registration, lifespan
│   ├── dashboard.py             Self-contained HTML dashboard (no template engine needed)
│   │
│   ├── api/
│   │   ├── __init__.py
│   │   ├── deps.py              Shared FastAPI dependencies (get_current_user)
│   │   └── routes/
│   │       ├── auth.py          POST /api/auth/register, POST /api/auth/login
│   │       ├── tasks.py         POST /api/tasks, GET /api/tasks, GET /api/tasks/:id
│   │       └── stats.py         GET /api/stats
│   │
│   ├── core/
│   │   ├── config.py            All env vars via pydantic-settings; @lru_cache singleton
│   │   ├── logging.py           structlog setup: JSON in prod, colored in dev
│   │   └── security.py          bcrypt hash/verify, JWT create/decode
│   │
│   ├── db/
│   │   ├── base.py              Async SQLAlchemy engine, AsyncSession, get_db() dependency
│   │   └── repository.py        Every DB query in one place (async for API, sync for worker)
│   │
│   ├── models/
│   │   ├── task.py              Task + TaskLog ORM models
│   │   ├── worker.py            Worker ORM model
│   │   └── user.py              User ORM model
│   │
│   ├── queue/
│   │   ├── client.py            Redis connection pool management (async + sync)
│   │   └── queue.py             enqueue(), dequeue_sync(), dead-letter push, depth check
│   │
│   ├── schemas/
│   │   ├── task.py              Pydantic schemas for task API (request + response shapes)
│   │   ├── auth.py              RegisterRequest, LoginRequest, TokenResponse
│   │   └── stats.py             StatsResponse, QueueDepths, WorkerInfo
│   │
│   └── worker/
│       ├── jobs/
│       │   └── handlers.py      One function per job type; HANDLERS dispatch dict
│       ├── webhook.py           HMAC-signed HTTP POST delivery
│       └── worker.py            Main worker: ThreadPoolExecutor, BRPOP loop, retry logic
│
├── alembic/
│   ├── env.py                   Alembic config: imports all models, sets DB URL
│   └── versions/
│       └── 001_initial_schema.py  Creates all tables and indexes
│
├── tests/
│   ├── unit/
│   │   ├── test_security.py     hash_password, verify_password, JWT encode/decode
│   │   ├── test_retry_logic.py  retry_delay_ms() math: base, growth, cap
│   │   └── test_job_handlers.py  All 5 handlers, execute_job dispatch, unknown type
│   └── integration/
│       └── test_api.py          Full API tests via httpx AsyncClient
│
├── scripts/
│   ├── seed.py                  Insert admin@jobqueue.dev / password123 for local dev
│   └── make_doc.py              Generate this .docx file
│
├── .github/
│   └── workflows/
│       └── ci.yml               4-stage GitHub Actions pipeline
│
├── locustfile.py                Load test: JobQueueUser + HealthCheckUser
├── alembic.ini                  Alembic config file (points to env.py)
├── pytest.ini                   asyncio_mode=auto, testpaths=tests
├── Dockerfile                   Single image for API + Worker
├── docker-compose.yml           4 services: postgres, redis, api, worker
├── requirements.txt             All Python dependencies pinned
└── .env                         Local environment variables (never commit this)""")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — ENVIRONMENT VARIABLES
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "7. Environment Variables — Configuration Guide")

env_vars = [
    ("DATABASE_URL",
     "postgresql+asyncpg://user:pass@host:5432/dbname",
     "Async database URL used by the FastAPI API. The 'postgresql+asyncpg://' prefix tells "
     "SQLAlchemy to use the asyncpg driver (required for async/await). "
     "Never use this in the worker — asyncpg requires an event loop."),
    ("DATABASE_URL_SYNC",
     "postgresql://user:pass@host:5432/dbname",
     "Sync database URL used by the worker process and Alembic. "
     "The 'postgresql://' prefix uses the psycopg2 driver (works in regular threads). "
     "Must point to the same database as DATABASE_URL."),
    ("DATABASE_POOL_SIZE",
     "10",
     "Number of persistent connections in the SQLAlchemy pool. "
     "For the worker, set this to at least WORKER_CONCURRENCY. "
     "Tuning this fixed the bottleneck found during Locust load testing."),
    ("DATABASE_MAX_OVERFLOW",
     "20",
     "Extra connections allowed beyond pool_size during traffic spikes. "
     "Total max connections = DATABASE_POOL_SIZE + DATABASE_MAX_OVERFLOW."),
    ("REDIS_URL",
     "redis://localhost:6379",
     "Redis connection URL. Both the API (async) and worker (sync) connect to Redis. "
     "In production, use AWS ElastiCache and set the URL to its endpoint."),
    ("SECRET_KEY",
     "change-this-to-a-long-random-string",
     "Key used to sign JWT tokens. Must be at least 32 random characters. "
     "If this leaks, attackers can forge valid tokens for any user. "
     "Generate with: python3 -c \"import secrets; print(secrets.token_hex(32))\""),
    ("WEBHOOK_SECRET",
     "change-this-too",
     "Key used to HMAC-sign webhook payloads. The receiver uses the same key "
     "to verify the signature. Prevents third parties from spoofing webhooks."),
    ("APP_ENV",
     "development or production",
     "Controls log format (colored console vs. JSON) and other environment-specific behavior. "
     "structlog switches to JSON output when APP_ENV=production."),
    ("WORKER_CONCURRENCY",
     "4",
     "Number of threads in the worker's ThreadPoolExecutor. "
     "Each thread independently processes one job at a time. "
     "Set based on your CPU cores and job type: IO-heavy jobs (email, API calls) "
     "can use more threads than CPU-heavy jobs (image processing)."),
    ("MAX_RETRIES",
     "3",
     "Maximum number of retry attempts before a task is moved to the dead-letter queue."),
    ("SENTRY_DSN",
     "https://xxx@oxxxxxxxx.ingest.sentry.io/yyy",
     "Optional. Sentry Data Source Name. If set, errors and performance traces are sent to Sentry. "
     "Leave empty in local development."),
]

for name, example, explanation in env_vars:
    body(doc, name, bold=True)
    code_block(doc, f"{name}={example}")
    body(doc, explanation, indent=1)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — RUNNING LOCALLY
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "8. Running Locally — Step-by-Step Setup")

subheading(doc, "Option A: Docker Compose (Recommended for Beginners)")
body(doc, "This is the simplest way. You only need Docker installed.")
code_block(doc, """\
# 1. Clone the repository
git clone https://github.com/yourname/job-queue
cd job-queue

# 2. Create .env from the example
cp .env.example .env
# Edit .env if needed — defaults work for local Docker setup

# 3. Build images and start all 4 services
docker compose up --build

# 4. Watch the logs — you should see:
#    postgres:  database system is ready to accept connections
#    redis:     Ready to accept connections
#    api:       Application startup complete
#    worker:    Worker a3f2b1c4 started with 4 threads

# 5. In a separate terminal, run the database seed:
docker compose exec api python scripts/seed.py""")
body(doc, "Now open:")
for url in [
    "http://localhost:8000/docs  — Swagger UI (test all endpoints here)",
    "http://localhost:8000/dashboard  — Live stats dashboard",
    "http://localhost:8000/health  — Health check (should return {status: ok})",
]:
    bullet(doc, url)

subheading(doc, "Option B: Manual Setup (Python 3.12)")
body(doc, "Use this if you want to run tests or develop without Docker.")
code_block(doc, """\
# 1. Install Python dependencies
pip install -r requirements.txt

# 2. Start PostgreSQL and Redis locally
# On macOS with Homebrew:
brew services start postgresql@16
brew services start redis

# On Ubuntu:
sudo systemctl start postgresql redis

# 3. Create the database
createdb job_queue   # or use psql

# 4. Set up .env
cp .env.example .env
# Set DATABASE_URL and DATABASE_URL_SYNC to use your local postgres user

# 5. Apply database migrations
alembic upgrade head
# Should print: Running upgrade -> 001, Create initial schema

# 6. Seed test user
python3 scripts/seed.py
# Prints: Created admin@jobqueue.dev

# 7. Start the API (terminal 1)
uvicorn app.main:app --reload --port 8000

# 8. Start the worker (terminal 2)
python3 -m app.worker.worker""")

note_box(doc, "COMMON ERROR:", (
    "If you see 'role postgres does not exist', your local PostgreSQL user "
    "is your system username (e.g. johndoe), not 'postgres'. "
    "Update DATABASE_URL_SYNC to: postgresql://johndoe@localhost:5432/job_queue"
))

subheading(doc, "Testing the API manually")
code_block(doc, """\
# Register a user
curl -X POST http://localhost:8000/api/auth/register \\
  -H "Content-Type: application/json" \\
  -d '{"email": "test@example.com", "password": "password123", "name": "Test"}'

# Login and save the token
TOKEN=$(curl -s -X POST http://localhost:8000/api/auth/login \\
  -H "Content-Type: application/json" \\
  -d '{"email": "test@example.com", "password": "password123"}' | python3 -c "import sys,json; print(json.load(sys.stdin)['access_token'])")

# Submit a task
curl -X POST http://localhost:8000/api/tasks \\
  -H "Authorization: Bearer $TOKEN" \\
  -H "Content-Type: application/json" \\
  -d '{"type": "email", "priority": "high", "payload": {"to": "a@b.com", "subject": "Hi"}}'

# Check the task status (replace TASK_ID with the id from the previous response)
curl http://localhost:8000/api/tasks/TASK_ID -H "Authorization: Bearer $TOKEN" """)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — RUNNING TESTS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "9. Running Tests — Unit, Integration, and Load Tests")

subheading(doc, "9.1  Unit Tests (fast, no dependencies)")
code_block(doc, """\
# Run just unit tests
pytest tests/unit/ -v

# Expected output:
# tests/unit/test_security.py::test_hash_password PASSED
# tests/unit/test_security.py::test_verify_password PASSED
# tests/unit/test_security.py::test_create_access_token PASSED
# tests/unit/test_security.py::test_decode_token PASSED
# tests/unit/test_security.py::test_token_has_expiry PASSED
# tests/unit/test_retry_logic.py::test_retry_delay_increases_exponentially PASSED
# tests/unit/test_retry_logic.py::test_retry_delay_first_attempt PASSED
# tests/unit/test_retry_logic.py::test_retry_delay_capped_at_1_hour PASSED
# tests/unit/test_retry_logic.py::test_retry_delay_base PASSED
# tests/unit/test_job_handlers.py::test_handle_email_success PASSED
# ... etc""")
body(doc, "What each unit test file covers:")
for item in [
    "test_security.py: hashing a password produces a non-reversible hash; verifying the correct password returns True; verifying the wrong password returns False; JWT tokens encode the correct payload; expired tokens are rejected",
    "test_retry_logic.py: attempt 0 = 5000ms, attempt 1 = 10000ms (doubles each time), very high attempt number is capped at 3,600,000ms (1 hour)",
    "test_job_handlers.py: each of the 5 handlers returns the expected result shape; missing required payload fields raise ValueError; execute_job dispatches to the right handler; unknown job type raises ValueError",
]:
    bullet(doc, item)

subheading(doc, "9.2  Integration Tests (require PostgreSQL + Redis)")
code_block(doc, """\
# Start dependencies first (if not using Docker):
# brew services start postgresql@16 redis

# Apply migrations to the test database
alembic upgrade head

# Run integration tests
pytest tests/integration/ -v

# Run full suite with coverage report
pytest --cov=app --cov-report=term-missing""")
body(doc, "What the integration tests cover:")
for item in [
    "GET /health returns 200 with status=ok",
    "POST /api/auth/register creates a user and returns a JWT",
    "POST /api/auth/login with correct credentials returns a JWT",
    "POST /api/auth/login with wrong password returns 401",
    "POST /api/tasks without auth returns 403",
    "POST /api/tasks with valid auth creates a task and returns 202",
    "GET /api/tasks/:id returns the correct task",
    "GET /api/tasks/:id for non-existent task returns 404",
    "POST /api/tasks with invalid type returns 422 with validation error details",
]:
    bullet(doc, item)

subheading(doc, "9.3  Load Tests (Locust)")
code_block(doc, """\
# Start the API first
uvicorn app.main:app --port 8000

# Run Locust with web UI
locust -f locustfile.py --host http://localhost:8000
# Open http://localhost:8089 in browser
# Set: Number of users = 100, Spawn rate = 10
# Click Start

# Or run headless (no browser):
locust -f locustfile.py --host http://localhost:8000 \\
  --users 100 --spawn-rate 10 --run-time 60s --headless""")
body(doc, (
    "The Locust file simulates realistic traffic: 60% of virtual users submit new tasks, "
    "30% check task status, 10% hit the stats endpoint. Each virtual user registers its own "
    "account on startup (simulating real user load, not shared credentials)."
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — CI/CD PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "10. CI/CD Pipeline — Automated Build and Deploy")

body(doc, (
    "CI/CD stands for Continuous Integration / Continuous Deployment. "
    "Every time code is pushed to the main branch (or a PR is opened), "
    "GitHub automatically runs tests, builds Docker images, and deploys to EC2. "
    "This means no manual deployments and no broken code reaches production."
))

subheading(doc, "10.1  Pipeline Stages")
code_block(doc, """\
Trigger: push to main or PR to main
         │
         ├─── Job 1: test-unit  (runs on every push/PR)
         │     runs-on: ubuntu-latest
         │     steps:
         │       - checkout code
         │       - setup Python 3.12
         │       - pip install -r requirements.txt
         │       - pytest tests/unit/ --cov=app --cov-report=xml
         │       - upload coverage to Codecov
         │
         ├─── Job 2: test-integration  (runs on every push/PR)
         │     services:
         │       postgres:16-alpine (with pg_isready health check)
         │       redis:7-alpine     (with redis-cli ping health check)
         │     steps:
         │       - checkout, setup Python, pip install
         │       - alembic upgrade head   ← applies migrations to the CI database
         │       - pytest tests/integration/ -v
         │
         ├─── Job 3: build-and-push  (only on push to main)
         │     needs: [test-unit, test-integration]  ← only runs if both pass
         │     steps:
         │       - docker buildx setup
         │       - docker login (using DOCKERHUB_USERNAME + DOCKERHUB_TOKEN secrets)
         │       - docker build + push API image:
         │           yourname/job-queue-api:latest
         │           yourname/job-queue-api:sha-abc1234
         │       - docker build + push Worker image:
         │           yourname/job-queue-worker:latest
         │           yourname/job-queue-worker:sha-abc1234
         │
         └─── Job 4: deploy  (only on push to main, after build-and-push)
               needs: build-and-push
               environment: production  ← requires manual approval (optional)
               steps:
                 - SSH into EC2 (using EC2_HOST + EC2_SSH_KEY secrets)
                 - cd ~/job-queue && docker compose pull
                 - docker compose up -d --no-build
                 - docker system prune -f   ← remove old images to save disk""")

subheading(doc, "10.2  Setting Up GitHub Secrets")
body(doc, "Go to GitHub repo → Settings → Secrets and variables → Actions → New repository secret:")
for secret in [
    "DOCKERHUB_USERNAME — your Docker Hub username",
    "DOCKERHUB_TOKEN — Docker Hub access token (generate at hub.docker.com → Account Settings → Security)",
    "EC2_HOST — your EC2 instance's public IP or DNS (e.g. ec2-54-123-45-67.compute.amazonaws.com)",
    "EC2_SSH_KEY — contents of your .pem private key file (paste the whole thing including BEGIN RSA PRIVATE KEY lines)",
]:
    bullet(doc, secret)

subheading(doc, "10.3  Why SHA-tagged Docker images?")
body(doc, (
    "Each push produces two tags: latest and sha-<commit-hash>. "
    "The sha tag lets you roll back to any previous version exactly. "
    "If a bad deployment causes errors, you run: docker compose down && docker pull yourname/job-queue-api:sha-abc1234 "
    "and docker compose up -d to instantly revert."
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — DEPLOYING TO AWS EC2
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "11. Deploying to AWS EC2 — Production Setup")

subheading(doc, "11.1  Launch an EC2 Instance")
for step in [
    "Log into AWS Console → EC2 → Launch Instance",
    "Choose Ubuntu Server 22.04 LTS (free tier eligible)",
    "Choose t3.small (2 vCPU, 2 GB RAM) — t2.micro is too small for 4 Docker services",
    "Create or select a key pair (.pem file) — you need this to SSH in",
    "Security group: allow inbound on port 22 (SSH), 8000 (API), and optionally 80/443",
    "Launch the instance",
]:
    numbered(doc, step)

subheading(doc, "11.2  Install Docker on EC2")
code_block(doc, """\
# SSH into EC2
ssh -i your-key.pem ubuntu@ec2-xx-xx-xx-xx.compute.amazonaws.com

# Update packages
sudo apt update && sudo apt upgrade -y

# Install Docker
sudo apt install -y docker.io docker-compose-v2

# Add ubuntu user to docker group (avoid using sudo with docker)
sudo usermod -aG docker ubuntu
newgrp docker

# Verify
docker --version
docker compose version""")

subheading(doc, "11.3  First Deploy")
code_block(doc, """\
# Clone your repo on EC2
git clone https://github.com/yourname/job-queue ~/job-queue
cd ~/job-queue

# Create .env with production values
nano .env
# Set DATABASE_URL, DATABASE_URL_SYNC, REDIS_URL, SECRET_KEY, WEBHOOK_SECRET
# For Docker Compose on EC2, the service hostnames are: postgres, redis
# DATABASE_URL=postgresql+asyncpg://postgres:strongpassword@postgres:5432/job_queue
# DATABASE_URL_SYNC=postgresql://postgres:strongpassword@postgres:5432/job_queue
# REDIS_URL=redis://redis:6379

# Start all services
docker compose up -d --build

# Apply migrations
docker compose exec api alembic upgrade head

# Create admin user
docker compose exec api python scripts/seed.py

# Check all services are running
docker compose ps

# Check API health
curl http://localhost:8000/health""")

subheading(doc, "11.4  Subsequent Deploys (handled by GitHub Actions)")
body(doc, (
    "After the first manual setup, all future deployments are fully automated. "
    "Push code to main → GitHub Actions runs tests → builds images → SSHes into EC2 → "
    "pulls new images → restarts containers. Zero manual work."
))

note_box(doc, "PRODUCTION RECOMMENDATION:", (
    "For a real production system, replace the docker-compose postgres and redis services "
    "with AWS RDS PostgreSQL and AWS ElastiCache Redis. "
    "These are managed services: AWS handles backups, failover, and scaling. "
    "Update your .env DATABASE_URL and REDIS_URL to point to the AWS endpoints."
), bg=YELLOW_BG)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — SECURITY
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "12. Security — Auth, Webhooks, Secrets")

subheading(doc, "12.1  Password Hashing with bcrypt")
body(doc, (
    "Passwords are never stored in plain text. When a user registers, we call "
    "hash_password(plain_password) which uses passlib's bcrypt implementation. "
    "bcrypt is intentionally slow (it runs thousands of internal iterations). "
    "This means even if your database is leaked, cracking passwords takes years per password. "
    "The hash includes a random salt, so two users with the same password have different hashes."
))
code_block(doc, """\
# In app/core/security.py
from passlib.context import CryptContext
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

def hash_password(plain: str) -> str:
    return pwd_context.hash(plain)         # returns '$2b$12$...' (bcrypt hash)

def verify_password(plain: str, hashed: str) -> bool:
    return pwd_context.verify(plain, hashed)   # constant-time comparison""")

subheading(doc, "12.2  JWT Authentication Flow")
code_block(doc, """\
# Login flow:
1. POST /api/auth/login { email, password }
2. Load user from DB by email
3. verify_password(plain, user.password_hash) → True/False
4. If True: create_access_token({sub: user.id, email: user.email})
   → signs with SECRET_KEY using HS256 algorithm
   → returns "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9..."
5. Client stores token in localStorage (web) or secure storage (mobile)

# Protected route flow:
1. Client sends: Authorization: Bearer eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...
2. FastAPI's Depends(get_current_user) extracts the token
3. decode_token(token) verifies the signature using SECRET_KEY
4. If valid and not expired: returns {id: ..., email: ...}
5. If invalid: raises HTTPException(403, "Invalid token")""")

subheading(doc, "12.3  Webhook HMAC-SHA256 Signing")
body(doc, (
    "When a task completes and there is a webhook_url, the worker fires a POST to that URL. "
    "But how does the receiver know the POST is really from our system and not a hacker? "
    "We sign the payload with HMAC-SHA256 and include the signature in a header."
))
code_block(doc, """\
# In app/worker/webhook.py
import hmac
import hashlib
import json

def deliver_webhook(task, settings):
    payload = json.dumps({
        "task_id": task.id,
        "status": task.status,
        "result": task.result,
    }).encode()

    # Compute signature
    sig = hmac.new(
        settings.WEBHOOK_SECRET.encode(),
        payload,
        hashlib.sha256
    ).hexdigest()

    headers = {
        "Content-Type": "application/json",
        "X-JobQueue-Signature": f"sha256={sig}"
    }
    # Send the POST request...

# Receiver verifies:
expected = hmac.new(secret, body, sha256).hexdigest()
received = request.headers["X-JobQueue-Signature"].replace("sha256=", "")
if hmac.compare_digest(expected, received):
    # Trust the payload
else:
    # Reject — not from our system""")
body(doc, (
    "Note: use hmac.compare_digest() not == for comparison. "
    "The == operator short-circuits (stops at the first different character), "
    "which creates a timing attack vector that could leak information about the correct signature. "
    "compare_digest always takes the same time regardless of where the strings differ."
))

subheading(doc, "12.4  Environment Variable Security")
for item in [
    "Never commit .env to git. Add it to .gitignore.",
    "In GitHub Actions, use Repository Secrets (never hardcode tokens in yaml)",
    "On EC2, set strict file permissions: chmod 600 .env",
    "Rotate SECRET_KEY if it is ever leaked (this invalidates all existing tokens)",
    "Use different SECRET_KEY values in development and production",
]:
    bullet(doc, item)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13 — OBSERVABILITY
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "13. Observability — Logging, Sentry, Dashboard")

subheading(doc, "13.1  Structured Logging with structlog")
body(doc, (
    "Traditional logs are strings. Structured logs are JSON objects. "
    "With JSON logs, you can filter and query them in log aggregation tools."
))
code_block(doc, """\
# Development output (colored, human-readable):
2026-07-06 10:23:45 [info     ] task_created    task_id=a3f2 type=email priority=high user_id=42
2026-07-06 10:23:46 [info     ] task_processing task_id=a3f2 thread=1 worker=b9c3
2026-07-06 10:23:46 [warning  ] task_retrying   task_id=a3f2 attempt=1 error=smtp_timeout delay_ms=10000
2026-07-06 10:23:56 [info     ] task_completed  task_id=a3f2 duration_ms=234

# Production output (JSON, machine-readable):
{"timestamp":"2026-07-06T10:23:45Z","level":"info","event":"task_created","task_id":"a3f2","type":"email"}
{"timestamp":"2026-07-06T10:23:46Z","level":"warning","event":"task_retrying","task_id":"a3f2","attempt":1}""")

subheading(doc, "13.2  Sentry Integration")
body(doc, "Sentry is configured in both app/main.py (API) and app/worker/worker.py (worker):")
code_block(doc, """\
import sentry_sdk
from sentry_sdk.integrations.fastapi import FastApiIntegration

if settings.SENTRY_DSN:
    sentry_sdk.init(
        dsn=settings.SENTRY_DSN,
        integrations=[FastApiIntegration()],
        traces_sample_rate=0.2,    # trace 20% of requests for performance data
        environment=settings.APP_ENV,
    )""")
body(doc, "What Sentry captures:")
for item in [
    "Unhandled exceptions in any route handler — with full stack trace and request context",
    "Worker job failures — with the job payload that triggered the error",
    "Performance traces — which DB query is slowest, which endpoint has the highest p95 latency",
    "Error frequency — how many times an error occurred in the last 24 hours",
]:
    bullet(doc, item)

subheading(doc, "13.3  Live Dashboard")
body(doc, (
    "Visit http://localhost:8000/dashboard for a self-refreshing HTML dashboard. "
    "It calls GET /api/stats every 3 seconds and renders:"
))
for item in [
    "Total tasks by status (pending, processing, completed, retrying, dead)",
    "Queue depths — how many items are currently in tasks:high, tasks:normal, tasks:low",
    "Registered workers — hostname, PID, status, tasks processed, tasks failed, last heartbeat",
    "A login form (token stored in localStorage, all API calls use the token)",
]:
    bullet(doc, item)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 14 — LOAD TESTING
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "14. Load Testing — Results and What We Fixed")

body(doc, (
    "Load testing means running the system under simulated heavy traffic to find performance problems "
    "before real users do. We used Locust to simulate 100 concurrent users."
))

subheading(doc, "14.1  The Test Setup")
code_block(doc, """\
locust -f locustfile.py --host http://localhost:8000 --users 100 --spawn-rate 10 --run-time 120s

User behavior (per virtual user):
  - on_start(): POST /api/auth/register (once)
  - @task(6): POST /api/tasks (60% of requests)
  - @task(3): GET  /api/tasks/:id (30% of requests)
  - @task(1): GET  /api/stats (10% of requests)""")

subheading(doc, "14.2  Results: Before the Fix")
body(doc, "Default SQLAlchemy pool_size=5:")
for m in [
    "0-50 users: p50=38ms, p95=52ms, error_rate=0%  — everything fine",
    "60 users:   p50=42ms, p95=2,400ms, error_rate=12%  — SPIKE",
    "70 users:   p50=55ms, p95=4,100ms, error_rate=23%  — getting worse",
]:
    bullet(doc, m)
body(doc, (
    "Error message in logs: 'QueuePool limit of size 5 overflow 10 reached, "
    "connection timed out, timeout 30'. "
    "The SQLAlchemy pool only had 5 persistent connections. At 60 concurrent requests, "
    "all 5 were in use. New requests waited up to 30 seconds for a connection to free up. "
    "After 30 seconds, they timed out with an error."
))

subheading(doc, "14.3  The Fix")
code_block(doc, """\
# In .env:
DATABASE_POOL_SIZE=20
DATABASE_MAX_OVERFLOW=40

# In app/core/config.py:
class Settings(BaseSettings):
    DATABASE_POOL_SIZE: int = 20
    DATABASE_MAX_OVERFLOW: int = 40

# In app/db/base.py:
engine = create_async_engine(
    settings.DATABASE_URL,
    pool_size=settings.DATABASE_POOL_SIZE,
    max_overflow=settings.DATABASE_MAX_OVERFLOW,
)""")

subheading(doc, "14.4  Results: After the Fix")
for m in [
    "100 users: p50=42ms, p95=95ms, error_rate=0%",
    "p95 improved from 2,400ms to 95ms — a 96% reduction",
    "Error rate went from 12% to 0%",
]:
    bullet(doc, m)
note_box(doc, "INTERVIEW TIP:", (
    "This is the most impressive part of the project to talk about in interviews. "
    "You found a real performance bottleneck using systematic load testing, "
    "diagnosed the root cause (connection pool exhaustion), and fixed it with a targeted config change. "
    "This shows you understand production systems, not just code."
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 15 — SCALING
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "15. Scaling the System — What Changes at 10x Load")

body(doc, (
    "The current setup handles ~500 tasks/minute comfortably on a single EC2 t3.small. "
    "Here is what you change as load grows:"
))

subheading(doc, "Scale the Worker (first bottleneck)")
body(doc, "Workers consume from the same Redis queues. Adding more workers increases throughput linearly.")
code_block(doc, """\
# In docker-compose.yml — increase replicas:
worker:
  deploy:
    replicas: 5   # was 1; now 5 workers, each with 4 threads = 20 concurrent jobs

# Or on EC2, run multiple worker containers:
docker compose up -d --scale worker=5""")

subheading(doc, "Scale the API (second bottleneck)")
body(doc, "Run multiple API instances behind a load balancer:")
code_block(doc, """\
# docker-compose.yml
api:
  deploy:
    replicas: 3   # 3 API instances on the same machine

# For cross-machine scaling: use AWS ALB in front of multiple EC2 instances
# The API is stateless (no in-memory state) — any instance can handle any request""")

subheading(doc, "Move to managed services (reliability)")
for item in [
    "Replace docker-compose postgres → AWS RDS PostgreSQL (Multi-AZ for failover, automated backups)",
    "Replace docker-compose redis → AWS ElastiCache Redis (managed, Redis Cluster for horizontal scale)",
    "Use AWS ECR instead of Docker Hub (private, IAM-controlled, lower latency from EC2)",
    "Add AWS CloudWatch for metrics and alarms (e.g., alert if DLQ depth > 100)",
]:
    bullet(doc, item)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 16 — COMMON MISTAKES
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "16. Common Mistakes and How to Avoid Them")

mistakes = [
    (
        "Committing .env to git",
        "Your SECRET_KEY and database passwords are in .env. "
        "If you push this to a public repo, attackers can access your database within minutes. "
        "Add .env to .gitignore immediately. Use .env.example (with fake values) as a template."
    ),
    (
        "Not handling worker crashes",
        "If the worker process crashes mid-job (after BRPOP, before completing), "
        "the task is left in 'processing' status forever. "
        "Add a recovery job that scans for tasks stuck in 'processing' for more than N minutes "
        "and re-enqueues them."
    ),
    (
        "Storing large payloads in Redis",
        "Redis is optimized for small messages (< 1MB). "
        "If you push a 10MB image to Redis, you will exhaust memory fast. "
        "Store large files in S3 and put only the S3 URL in the task payload."
    ),
    (
        "Mixing async and sync SQLAlchemy",
        "asyncpg (async) and psycopg2 (sync) cannot be used interchangeably. "
        "The API uses asyncpg. The worker uses psycopg2. "
        "Using the wrong driver causes cryptic errors about event loops or blocking calls. "
        "This project has two separate DATABASE_URL variables to keep them separate."
    ),
    (
        "Not rate-limiting retries",
        "If you retry immediately on failure, a broken downstream service gets hammered. "
        "Always use exponential backoff. The cap (3,600,000ms = 1 hour) prevents retries "
        "from waiting so long they are never useful."
    ),
    (
        "Forgetting the DLQ",
        "Without a DLQ, failed tasks disappear. "
        "You have no way to know how many tasks failed or why. "
        "Always route max-retry failures to a DLQ and set up monitoring on its depth."
    ),
    (
        "Running pytest without asyncio_mode=auto",
        "FastAPI routes are async functions. pytest needs pytest-asyncio with asyncio_mode=auto "
        "to run async test functions. Without it, async tests either fail or silently pass without executing. "
        "Check pytest.ini has: asyncio_mode = auto"
    ),
]
for mistake, fix in mistakes:
    body(doc, mistake, bold=True)
    body(doc, fix, indent=1)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 17 — INTERVIEW Q&A
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "17. Interview Questions & Detailed Answers")
note_box(doc, "STRATEGY:", (
    "When answering system design questions, use this structure: "
    "(1) Define the concept in one sentence. "
    "(2) Explain how it is implemented in THIS project specifically. "
    "(3) Mention the trade-off or design decision. "
    "This shows you understand both theory and practice."
))

qa_pairs = [
    (
        "What is a job queue and why would you use one?",
        "A job queue is a system that decouples slow or resource-intensive work from the HTTP request cycle. "
        "The API accepts the request, saves the task, and returns 202 Accepted immediately. "
        "A background worker picks up the task and processes it asynchronously. "
        "You use it when work is too slow for a synchronous HTTP response (email takes 1-3s), "
        "when you need retries (email servers go down), or when you need to scale processing "
        "independently from the API (add more workers without adding API instances)."
    ),
    (
        "Walk me through how a task moves from POST /api/tasks to completion.",
        "Five steps. (1) The API validates the JWT, validates the request with Pydantic, "
        "inserts a row into PostgreSQL with status=pending, and returns 202. "
        "(2) The API calls LPUSH on the appropriate Redis list (tasks:high, tasks:normal, or tasks:low). "
        "(3) A worker thread wakes up from BRPOP, receives the task ID. "
        "(4) The worker loads the full task from PostgreSQL, updates status=processing, "
        "and calls execute_job(task.type, task.payload). "
        "(5) On success: updates status=completed and fires the webhook. "
        "On failure: increments retry_count, calculates backoff delay, re-enqueues after the delay. "
        "After 3 failures: moves to the dead-letter queue."
    ),
    (
        "How does priority queuing work? Why three Redis lists instead of one sorted set?",
        "We have three Redis lists: tasks:high, tasks:normal, tasks:low. "
        "Producers LPUSH to the appropriate list. Workers call BRPOP on all three in order. "
        "Redis checks lists left-to-right: if tasks:high is non-empty, it always returns from there first. "
        "A Redis Sorted Set (ZSET) would also work — you'd use the priority as the score and ZPOPMIN. "
        "We chose lists because BRPOP is atomic across multiple lists, blocking is built-in, "
        "and the logic is simpler with just three priority levels. "
        "ZSET would be better if you needed numeric priority scores (1-100) instead of three levels."
    ),
    (
        "What is exponential backoff and why is the delay capped?",
        "Exponential backoff means each retry waits longer than the previous one: "
        "delay = min(5000ms * 2^attempt, 3,600,000ms). "
        "Attempt 1 = 10s, attempt 2 = 20s, attempt 3 = 40s. "
        "The motivation: if the downstream service (email server, payment API) is overloaded, "
        "retrying immediately makes it worse and wastes resources. "
        "Backing off gives it time to recover. "
        "The cap at 3,600,000ms (1 hour) prevents retries from waiting so long they become useless "
        "— you want them to eventually run, not wait 136 years on attempt 50."
    ),
    (
        "Explain the threading model in the worker. Why threads and not asyncio?",
        "The worker uses Python's ThreadPoolExecutor with N threads (configurable via WORKER_CONCURRENCY). "
        "Each thread runs its own independent loop: BRPOP, execute job, repeat. "
        "We use threads instead of asyncio because BRPOP is a blocking network call — "
        "it blocks the calling thread for up to 5 seconds if the queue is empty. "
        "In an asyncio event loop, one blocking call would stall ALL coroutines. "
        "Threads handle blocking calls naturally: each thread sleeps in the kernel during BRPOP, "
        "the OS wakes it when Redis has data, other threads continue unaffected. "
        "Thread safety: each thread creates its own SQLAlchemy Session and closes it in finally, "
        "so there is no shared mutable state between threads."
    ),
    (
        "How does graceful shutdown work?",
        "We use threading.Event called _shutdown. "
        "The main thread calls _shutdown.wait() which blocks forever. "
        "When SIGTERM or SIGINT arrives (docker stop, Ctrl-C, EC2 scale-in event), "
        "the signal handler calls _shutdown.set(). "
        "Worker threads check 'while not _shutdown.is_set()' at the top of each loop. "
        "When _shutdown is set, each thread finishes its current job and exits normally. "
        "The ThreadPoolExecutor context manager waits for all threads to finish before the process exits. "
        "This means no job is abandoned mid-execution during a rolling deploy or EC2 shutdown."
    ),
    (
        "What is HMAC and why do you sign webhook payloads?",
        "HMAC (Hash-based Message Authentication Code) is a way to sign data using a shared secret key. "
        "We compute HMAC-SHA256 of the webhook payload using WEBHOOK_SECRET and put it in the "
        "X-JobQueue-Signature header. The receiver does the same computation and compares. "
        "If they match, the payload is authentic and unmodified. "
        "Without this, anyone who knows your webhook URL can send fake completion events. "
        "We use hmac.compare_digest() for comparison to prevent timing attacks: "
        "regular == exits early on the first differing character, leaking information about "
        "the correct signature through response time differences."
    ),
    (
        "Why does the API use asyncpg but the worker uses psycopg2?",
        "FastAPI is an async framework — all route handlers are async functions running on an asyncio event loop. "
        "asyncpg is the async PostgreSQL driver: awaiting a query yields control to the event loop "
        "so other requests can be handled while waiting for the DB. "
        "The worker uses regular threads, not asyncio. psycopg2 is the sync driver: "
        "it blocks the calling thread while the query runs, which is fine in a thread. "
        "If you used asyncpg in the worker, it would complain about 'no running event loop'. "
        "If you used psycopg2 in the async API, awaiting queries would block the event loop, "
        "making the API handle one request at a time."
    ),
    (
        "What is Alembic and why not just use SQLAlchemy's create_all()?",
        "Alembic is a database migration tool. SQLAlchemy's Base.metadata.create_all() creates all tables "
        "if they don't exist, but it never modifies existing tables. "
        "If you add a column to a model, create_all() won't add it to the existing table. "
        "Alembic tracks which migrations have been applied (in a alembic_version table) "
        "and runs only the new ones. In CI we run alembic upgrade head before integration tests "
        "to ensure the schema is up to date. In production, the same command runs before starting "
        "the API container. Each migration also has a downgrade() function to reverse the change."
    ),
    (
        "Walk me through the CI/CD pipeline and what each stage does.",
        "GitHub Actions with 4 stages. "
        "Stage 1 (test-unit): installs Python dependencies, runs pytest tests/unit/. "
        "No services needed — fast, runs in ~30 seconds. "
        "Stage 2 (test-integration): GitHub Actions spins up postgres:16 and redis:7 as Docker service containers "
        "with health checks. We run alembic upgrade head, then pytest tests/integration/. "
        "This stage catches bugs that only appear with real database/Redis interactions. "
        "Stage 3 (build-and-push): runs only on merge to main, after both test stages pass. "
        "Builds Docker images for API and Worker using Buildx (multi-platform support), "
        "pushes to Docker Hub tagged with latest and the git SHA. "
        "Stage 4 (deploy): SSHes into EC2 using a stored private key, pulls new images, "
        "restarts containers with docker compose up -d, prunes old images to free disk space."
    ),
    (
        "What did the Locust load test reveal and how did you fix it?",
        "At ~60 concurrent users, POST /api/tasks p95 latency jumped from 52ms to 2,400ms "
        "and error rate hit 12%. The structlog logs showed 'QueuePool limit of size 5 overflow 10 reached, "
        "connection timed out'. SQLAlchemy's default connection pool has 5 persistent connections "
        "and 10 overflow slots. With 60 concurrent requests each needing a DB connection, "
        "we exhausted the pool. New requests waited 30 seconds for a connection, then timed out. "
        "The fix: set DATABASE_POOL_SIZE=20 and DATABASE_MAX_OVERFLOW=40 in .env. "
        "After the fix, under 100 users: p95=95ms, error rate=0%."
    ),
    (
        "How would you make tasks idempotent?",
        "Idempotency means running a task twice produces the same result as once. "
        "For example, sending the same welcome email twice is bad. "
        "Approach: before executing a job, check if a result already exists for this task ID. "
        "If it does, skip execution and return the cached result. "
        "For email specifically: store the message_id returned by the email provider "
        "and check if a record with that recipient+subject+date already has a message_id. "
        "Alternatively, for any job: use a Redis key 'job:executed:{task_id}' with a TTL. "
        "If the key exists, skip execution. Set it on first completion. "
        "This handles at-least-once delivery without double-processing."
    ),
    (
        "How would you handle tasks that need to run on a schedule?",
        "This project handles one-off tasks. For scheduled tasks (run daily, run every hour): "
        "Option 1: Use a cron job (crontab on EC2) that calls POST /api/tasks at scheduled times. "
        "Simple but external to the codebase. "
        "Option 2: Add a scheduler service using APScheduler or Celery Beat "
        "that generates tasks on a schedule and enqueues them. "
        "Option 3: For AWS, use EventBridge (CloudWatch Events) to trigger a Lambda "
        "that calls POST /api/tasks on your EC2 API. Fully managed, no extra process."
    ),
    (
        "What is the difference between a task queue and a message broker?",
        "A message queue is a general-purpose publish-subscribe system: producers send messages, "
        "consumers receive them. Examples: RabbitMQ, AWS SQS, Apache Kafka. "
        "A task queue is built on top of a message broker specifically for executing code (jobs). "
        "It adds: retry logic, result storage, task status tracking, priority, scheduling. "
        "In this project, Redis lists are the message broker and the Python code around them "
        "implements the task queue semantics (retries, DLQ, status updates). "
        "Celery is a more feature-complete task queue library that can use Redis or RabbitMQ as its broker."
    ),
    (
        "What is Pydantic and how does it help with validation?",
        "Pydantic is a Python data validation library. You define a class with type annotations "
        "and Pydantic automatically validates incoming data against those types. "
        "In FastAPI, when you declare a route parameter as a Pydantic model, FastAPI automatically "
        "parses the JSON body, validates it, and returns a 422 error with field-level details if validation fails. "
        "In this project, TaskCreate validates that 'type' is one of the allowed enum values, "
        "'priority' is high/normal/low, and 'payload' is a dict. "
        "Without Pydantic, you would write if/else validation code in every route handler."
    ),
    (
        "Why 202 Accepted and not 200 OK for task creation?",
        "HTTP status codes have specific meanings. "
        "200 OK means the request was fulfilled completely and synchronously. "
        "202 Accepted means 'we received the request and will process it, but have not done so yet'. "
        "This is the correct code for async operations. "
        "It signals to clients that they need to poll (GET /api/tasks/:id) to check completion. "
        "201 Created would also be wrong because 201 implies the resource is fully created and ready, "
        "not just queued. Using 202 is an important API design detail that shows you understand REST semantics."
    ),
    (
        "How does the system ensure a task is never silently lost?",
        "Three layers of durability: "
        "(1) Write to PostgreSQL first (status=pending) before pushing to Redis. "
        "If Redis is down, the task record exists and can be recovered. "
        "(2) The worker loads task details from PostgreSQL (not from the Redis message). "
        "The Redis message only contains the task ID. "
        "(3) Dead-letter queue: tasks that fail all retries go to tasks:dead in Redis "
        "and status=dead in PostgreSQL, never silently dropped. "
        "Gap: if the worker pops from Redis but crashes before updating PostgreSQL, "
        "the task is stuck in 'processing'. A recovery job scanning for tasks in "
        "'processing' for >10 minutes and re-enqueueing them would close this gap."
    ),
    (
        "How do you monitor this system in production?",
        "Three tools working together: "
        "(1) structlog: every event (task created, started, completed, failed) is logged as JSON. "
        "In production these flow to CloudWatch Logs where you can set metric filters and alarms. "
        "(2) Sentry: captures every unhandled exception in the API and worker in real time, "
        "sends email/Slack alerts, and shows the full stack trace. "
        "(3) The /api/stats endpoint + dashboard: shows queue depths (how many tasks in each queue), "
        "task counts by status, and worker heartbeats. "
        "If queue depth keeps growing, it means workers are falling behind — time to scale up. "
        "If DLQ depth is non-zero, something is broken — time to investigate."
    ),
    (
        "What is a connection pool and why does it matter?",
        "A connection pool is a set of pre-established database connections that are reused across requests. "
        "Opening a new PostgreSQL connection takes ~50-100ms (TCP handshake, authentication, startup). "
        "With a pool, connections are established at startup and reused. "
        "A request borrows a connection from the pool, runs its query, and returns the connection. "
        "pool_size controls how many persistent connections are kept open. "
        "max_overflow allows temporary extra connections during traffic spikes. "
        "If all connections are in use and the overflow is full, new requests wait. "
        "This is exactly the bottleneck Locust revealed in our load test — "
        "the default pool_size=5 was exhausted at 60 concurrent users."
    ),
    (
        "Why use Docker for deployment instead of just copying files?",
        "Docker solves the 'works on my machine' problem. "
        "It packages the code, Python version, all pip packages, and system libraries "
        "(like libpq for psycopg2) into an image. "
        "The image runs identically on your laptop, in CI, and on EC2. "
        "Without Docker, you would need to manually install Python 3.12, pip packages, "
        "and system dependencies on each server. Different server setups lead to different behavior. "
        "Docker also enables easy horizontal scaling (start more containers), "
        "rollbacks (pull the previous SHA-tagged image), "
        "and resource isolation (each container gets its own filesystem and process namespace)."
    ),
]

for q, a in qa_pairs:
    qa_block(doc, q, a)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 18 — GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "18. Glossary")

terms = [
    ("ACID", "Atomicity, Consistency, Isolation, Durability — properties of reliable database transactions. A task insert either fully succeeds or fully fails; no partial writes."),
    ("asyncpg", "An async PostgreSQL driver for Python. Used in the FastAPI API. Requires an asyncio event loop."),
    ("BRPOP", "Blocking Right POP — Redis command that blocks until an item is available in one of the specified lists, then returns and removes it."),
    ("CI/CD", "Continuous Integration / Continuous Deployment. CI = auto-run tests on every commit. CD = auto-deploy passing builds to production."),
    ("Dead-Letter Queue (DLQ)", "A separate queue where permanently failed messages/tasks are sent after all retries are exhausted. Used for manual inspection and recovery."),
    ("Docker", "A tool that packages an app and its dependencies into a portable image. Run anywhere with identical behavior."),
    ("Exponential Backoff", "A retry strategy where each retry waits twice as long as the previous one. Prevents hammering a struggling service."),
    ("FastAPI", "A modern Python web framework for building JSON APIs. Built on Starlette (async HTTP) and Pydantic (validation)."),
    ("HMAC", "Hash-based Message Authentication Code. A cryptographic signature computed with a shared secret key. Used to verify that a message was not tampered with."),
    ("Idempotent", "An operation that produces the same result when applied multiple times. Important for retry safety."),
    ("JWT", "JSON Web Token. A base64-encoded, signed token used for stateless authentication. Contains: header.payload.signature."),
    ("LPUSH", "Redis command: Left PUSH — adds an item to the front (left end) of a list."),
    ("ORM", "Object-Relational Mapper. Maps Python classes to database tables. SQLAlchemy is the ORM in this project."),
    ("Pydantic", "Python data validation library. Defines data shapes as Python classes with type annotations. FastAPI uses it for request/response validation."),
    ("p50 / p95 / p99", "Percentile latencies. p95=95ms means 95% of requests completed in 95ms or less. p50 is the median. p99 is the worst 1%."),
    ("psycopg2", "A sync PostgreSQL driver for Python. Used in the worker process. Blocks the calling thread during queries."),
    ("Redis", "An in-memory data structure server. Used as the message broker in this project. Extremely fast: sub-millisecond operations."),
    ("SQLAlchemy", "Python ORM and SQL toolkit. Version 2.0 introduced a clean async API. Used for all PostgreSQL access."),
    ("structlog", "Python structured logging library. Produces machine-parseable JSON logs in production."),
    ("ThreadPoolExecutor", "Python standard library class that manages a pool of threads. Workers use it to process multiple jobs concurrently."),
    ("202 Accepted", "HTTP status code meaning: the request was received and will be processed asynchronously. Appropriate for queued operations."),
    ("422 Unprocessable Entity", "HTTP status code meaning: the request body failed validation. FastAPI returns this automatically when a Pydantic model validation fails."),
]

for term, definition in terms:
    body(doc, term, bold=True)
    body(doc, definition, indent=1)

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = os.path.expanduser("~/Downloads/Job_Queue_Complete_Guide.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
